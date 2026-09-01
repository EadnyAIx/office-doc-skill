"""多格式文档解析器：统一接口解析 PDF / Word / Markdown / TXT。"""

from pathlib import Path
from typing import Dict, List


class ParseError(Exception):
    """解析异常。"""


def parse_document(path: str) -> Dict:
    """解析文档，返回结构化内容。

    Args:
        path: 文档路径

    Returns:
        dict: {"type": 文档类型, "title": 标题, "content": 正文,
               "tables": 表格列表, "path": 路径}
    """
    file_path = Path(path)
    if not file_path.exists():
        raise ParseError(f"文件不存在: {path}")
    if not file_path.is_file():
        raise ParseError(f"不是文件: {path}")

    ext = file_path.suffix.lower()
    parser = _get_parser(ext)
    return parser(file_path)


def _get_parser(ext: str):
    """根据扩展名选择解析器。"""
    if ext == ".pdf":
        return _parse_pdf
    if ext in (".docx", ".doc"):
        return _parse_docx
    if ext in (".md", ".markdown"):
        return _parse_markdown
    if ext == ".txt":
        return _parse_txt
    raise ParseError(f"不支持的文件格式: {ext}（支持 pdf/docx/md/txt）")


def _parse_pdf(path: Path) -> Dict:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ParseError("未安装 pypdf，请运行 pip install pypdf")

    try:
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)

        content = "\n\n".join(pages)
        return {
            "type": "pdf",
            "title": path.stem,
            "content": content.strip(),
            "tables": [],
            "path": str(path),
            "meta": {"pages": len(reader.pages)},
        }
    except Exception as e:
        raise ParseError(f"PDF 解析失败: {e}")


def _parse_docx(path: Path) -> Dict:
    try:
        import docx
    except ImportError:
        raise ParseError("未安装 python-docx，请运行 pip install python-docx")

    try:
        document = docx.Document(str(path))

        # 提取标题（段落样式以 Heading 开头）
        title = path.stem
        paragraphs = []
        tables = []

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if title == path.stem and para.style.name.startswith("Heading"):
                title = text
            paragraphs.append(text)

        # 提取表格
        for table in document.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                tables.append(rows)

        return {
            "type": "docx",
            "title": title,
            "content": "\n".join(paragraphs),
            "tables": tables,
            "path": str(path),
            "meta": {"paragraphs": len(paragraphs), "tables": len(tables)},
        }
    except Exception as e:
        raise ParseError(f"Word 解析失败: {e}")


def _parse_markdown(path: Path) -> Dict:
    text = path.read_text(encoding="utf-8-sig", errors="ignore")
    lines = text.splitlines()

    # 提取标题（第一个 # 开头的行）
    title = path.stem
    for line in lines:
        if line.startswith("#"):
            title = line.lstrip("#").strip()
            break

    # 提取表格（| 分隔的行）
    tables = []
    i = 0
    while i < len(lines):
        if lines[i].strip().startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            header = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
                j += 1
            tables.append({"header": header, "rows": rows})
            i = j
        else:
            i += 1

    return {
        "type": "markdown",
        "title": title,
        "content": text,
        "tables": tables,
        "path": str(path),
        "meta": {"lines": len(lines), "tables": len(tables)},
    }


def _parse_txt(path: Path) -> Dict:
    text = path.read_text(encoding="utf-8-sig", errors="ignore")
    return {
        "type": "txt",
        "title": path.stem,
        "content": text.strip(),
        "tables": [],
        "path": str(path),
        "meta": {"chars": len(text)},
    }


def format_plain(parsed: Dict) -> str:
    """将解析结果格式化为易读文本。"""
    lines = [f"📄 文档: {parsed['title']} ({parsed['type'].upper()})"]
    lines.append("=" * 50)
    if parsed["content"]:
        lines.append(parsed["content"])
    for t in parsed["tables"]:
        if isinstance(t, dict):
            lines.append("\n📊 表格:")
            lines.append(" | ".join(t["header"]))
            for row in t["rows"][:10]:
                lines.append(" | ".join(row))
        else:
            lines.append("\n📊 表格:")
            for row in t[:10]:
                lines.append(" | ".join(row))
    return "\n".join(lines)


def format_json(parsed: Dict) -> str:
    """将解析结果格式化为 JSON 字符串。"""
    import json
    # 序列化时处理表格等结构
    def clean(obj):
        if isinstance(obj, dict):
            return {k: clean(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [clean(i) for i in obj]
        return obj
    return json.dumps(clean(parsed), ensure_ascii=False, indent=2)
