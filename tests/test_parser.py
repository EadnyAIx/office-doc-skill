"""parser 单元测试。"""

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from parser import parse_document, ParseError


def _write(path: Path, content: str, encoding="utf-8"):
    path.write_text(content, encoding=encoding)
    return str(path)


def test_parse_txt():
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "test.txt"
        result = parse_document(_write(p, "这是第一行\n这是第二行"))
        assert result["type"] == "txt"
        assert "第一行" in result["content"]


def test_parse_markdown():
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "note.md"
        content = "# 标题\n\n正文内容\n\n| 列A | 列B |\n|---|---|\n| 1 | 2 |"
        result = parse_document(_write(p, content))
        assert result["type"] == "markdown"
        assert result["title"] == "标题"
        assert "正文内容" in result["content"]
        assert len(result["tables"]) == 1


def test_parse_missing_file():
    try:
        parse_document("E:/nonexistent_file_xyz.pdf")
        assert False, "应抛出异常"
    except ParseError as e:
        assert "不存在" in str(e)


def test_parse_unsupported_ext():
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "test.xyz"
        p.write_text("x")
        try:
            parse_document(str(p))
            assert False, "应抛出异常"
        except ParseError as e:
            assert "不支持" in str(e)


def test_parse_docx():
    try:
        import docx  # noqa
    except ImportError:
        return  # 未安装则跳过
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "test.docx"
        document = docx.Document()
        document.add_heading("Word标题", level=1)
        document.add_paragraph("Word正文内容")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        document.save(str(p))
        result = parse_document(str(p))
        assert result["type"] == "docx"
        assert "Word正文内容" in result["content"]
        assert len(result["tables"]) == 1
