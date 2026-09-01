"""格式转换器：Markdown↔HTML、TXT→Word、Markdown→Word。"""

from pathlib import Path
from typing import List

from parser import parse_document


class ConvertError(Exception):
    """转换异常。"""


SUPPORTED = {"html", "docx", "txt"}


def convert_file(src_path: str, to_format: str, out_path: str = None) -> str:
    """转换文档格式。

    Args:
        src_path: 源文件路径
        to_format: 目标格式 (html / docx / txt)
        out_path: 输出路径（默认 output/ 目录）

    Returns:
        输出文件路径
    """
    src = Path(src_path)
    if not src.exists():
        raise ConvertError(f"文件不存在: {src_path}")

    to_format = to_format.lower().lstrip(".")
    if to_format not in SUPPORTED:
        raise ConvertError(f"不支持的目标格式: {to_format}（支持 html/docx/txt）")

    out = _resolve_out_path(src, to_format, out_path)

    ext = src.suffix.lower()
    if to_format == "html":
        if ext in (".md", ".markdown", ".txt"):
            content = _md_or_txt_to_html(src)
            out.write_text(content, encoding="utf-8")
        elif ext == ".docx":
            # 先解析再转 markdown 再转 html
            parsed = parse_document(str(src))
            md = _parsed_to_md(parsed)
            out.write_text(_md_to_html(md), encoding="utf-8")
        else:
            raise ConvertError(f"不支持 {ext} → html")
    elif to_format == "docx":
        if ext in (".md", ".markdown", ".txt"):
            text = src.read_text(encoding="utf-8-sig", errors="ignore")
            _text_to_docx(text, out)
        else:
            raise ConvertError(f"不支持 {ext} → docx（请先转成 md/txt）")
    elif to_format == "txt":
        if ext == ".docx":
            parsed = parse_document(str(src))
            out.write_text(parsed["content"], encoding="utf-8")
        elif ext == ".pdf":
            parsed = parse_document(str(src))
            out.write_text(parsed["content"], encoding="utf-8")
        else:
            raise ConvertError(f"不支持 {ext} → txt")

    return str(out)


def _resolve_out_path(src: Path, to_format: str, out_path: str = None) -> Path:
    """解析输出路径。"""
    if out_path:
        out = Path(out_path)
    else:
        out_dir = src.parent / "output"
        out_dir.mkdir(parents=True, exist_ok=True)
        out = out_dir / f"{src.stem}.{to_format}"
    out.parent.mkdir(parents=True, exist_ok=True)
    return out


def _md_to_html(md_text: str) -> str:
    """Markdown 转 HTML。"""
    import markdown
    body = markdown.markdown(
        md_text,
        extensions=["tables", "fenced_code", "codehilite"],
    )
    return f"<!DOCTYPE html>\n<html>\n<head>\n<meta charset='utf-8'>\n<title>Converted</title>\n<style>body{{max-width:800px;margin:auto;padding:2em;font-family:sans-serif;line-height:1.6}}table{{border-collapse:collapse}}td,th{{border:1px solid #ccc;padding:6px}}code{{background:#f4f4f4;padding:2px 4px}}</style>\n</head>\n<body>\n{body}\n</body>\n</html>"


def _md_or_txt_to_html(src: Path) -> str:
    """读取 md/txt 并转 HTML。"""
    text = src.read_text(encoding="utf-8-sig", errors="ignore")
    if src.suffix.lower() in (".txt",):
        # 简单转义为段落
        from html import escape
        paras = [escape(p.strip()) for p in text.split("\n\n") if p.strip()]
        body = "\n".join(f"<p>{p}</p>" for p in paras)
        return f"<!DOCTYPE html>\n<html><head><meta charset='utf-8'><title>{src.stem}</title></head><body>{body}</body></html>"
    return _md_to_html(text)


def _parsed_to_md(parsed: dict) -> str:
    """解析结果转 Markdown。"""
    parts = [f"# {parsed['title']}", "", parsed["content"]]
    for t in parsed["tables"]:
        if isinstance(t, dict):
            parts.append("")
            parts.append("| " + " | ".join(t["header"]) + " |")
            parts.append("|" + "|".join(["---"] * len(t["header"])) + "|")
            for row in t["rows"]:
                parts.append("| " + " | ".join(row) + " |")
    return "\n".join(parts)


def _text_to_docx(text: str, out: Path) -> None:
    """纯文本转 Word 文档。"""
    try:
        import docx
    except ImportError:
        raise ConvertError("未安装 python-docx，请运行 pip install python-docx")

    document = docx.Document()
    # 简单解析：空行分段，以 # 开头为标题
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            document.add_heading(block[2:].strip(), level=1)
        elif block.startswith("## "):
            document.add_heading(block[3:].strip(), level=2)
        else:
            for line in block.split("\n"):
                if line.strip().startswith("- "):
                    document.add_paragraph(line.strip()[2:], style="List Bullet")
                else:
                    document.add_paragraph(line.strip())
    document.save(str(out))
